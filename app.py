from flask import Flask, request, render_template, send_file
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import os

app = Flask(__name__)

# Function to extract headings and subheadings with placeholders for page numbers
def extract_headings_and_subheadings_with_page_numbers(docx_file):
    doc = Document(docx_file)
    index_data = []
    hnum = 0
    shnum = 0.0
    virtual_page_num = 1  # Virtual page tracker for index only

    for para in doc.paragraphs:
        text = para.text.strip()

        if "<m>" in text and "</m>" in text:
            hnum += 1
            shnum = hnum
            start = text.find("<m>") + 3
            end = text.find("</m>")
            heading_text = f"{text[start:end]}"
            index_data.append((hnum, heading_text, virtual_page_num))

        elif "<s>" in text and "</s>" in text:
            shnum += 0.1
            start = text.find("<s>") + 3
            end = text.find("</s>")
            subheading_text = f"\t{round(shnum, 1)} {text[start:end]}"
            index_data.append(("", subheading_text, virtual_page_num))

        # Increment virtual page number when <b></b> is encountered
        if "<b></b>" in text:
            virtual_page_num += 1

    return index_data

# Function to create PDF with index on the first page and content on subsequent pages
def create_pdf_with_index_and_content(docx_file, pdf_file, index_data):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_file, pagesize=A4)
    elements = []

    # Generate the index on the first page
    elements.append(Paragraph("Index", styles["Title"]))
    elements.append(Spacer(1, 12))

    # Create the index table data with virtual page numbers
    table_data = [["Ch. No", "Title", "Page Number"]]
    for hnum, title, page_num in index_data:
        if hnum:  # Main heading
            table_data.append([hnum, Paragraph(f"<b>{title}</b>", styles["Normal"]), page_num])
        else:  # Subheading
            table_data.append(["", Paragraph(f"&bull; {title}", styles["Normal"]), page_num])

    # Create the table for index with boxed layout
    table = Table(table_data, colWidths=[1 * inch, 4 * inch, 1 * inch])
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONT', (0, 1), (-1, -1), 'Helvetica')
    ]))
    elements.append(table)

    # Add a page break to start the main content on a new page
    elements.append(PageBreak())

    # Add document content without actual page breaks for <b></b>
    doc_content = Document(docx_file)
    for para in doc_content.paragraphs:
        text = para.text.strip()
        if "<m>" in text and "</m>" in text:
            # Main heading
            start = text.find("<m>") + 3
            end = text.find("</m>")
            heading_text = text[start:end]
            elements.append(Paragraph(f"<b>{heading_text}</b>", styles["Heading1"]))
            elements.append(Spacer(1, 12))
        elif "<s>" in text and "</s>" in text:
            # Subheading
            start = text.find("<s>") + 3
            end = text.find("</s>")
            subheading_text = text[start:end]
            elements.append(Paragraph(f"<b>{subheading_text}</b>", styles["Heading2"]))
            elements.append(Spacer(1, 8))
        else:
            # Body text with wrapping
            elements.append(Paragraph(text, styles["BodyText"]))
            elements.append(Spacer(1, 6))

    # Build the PDF with the final content
    doc.build(elements)

# Route for file upload page
@app.route('/')
def upload_file():
    return render_template('index.html')

# Route for file conversion
@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return "No file part"
    file = request.files['file']
    
    if file.filename == '':
        return "No selected file"
    
    docx_file_path = os.path.join('uploads', file.filename)
    file.save(docx_file_path)
    
    pdf_file_path = os.path.join('uploads', file.filename.replace('.docx', '.pdf'))
    
    # Extract headings and subheadings for index creation with virtual page numbering
    index_data = extract_headings_and_subheadings_with_page_numbers(docx_file_path)
    
    # Create PDF with index on the first page and document content
    create_pdf_with_index_and_content(docx_file_path, pdf_file_path, index_data)

    return send_file(pdf_file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
